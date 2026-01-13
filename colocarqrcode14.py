from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def substituir_texto_por_imagem(docx_path, texto, imagem_path, output_path):
    """
    Substitui o texto especificado por uma imagem dentro de uma tabela em um documento do Word.

    :param docx_path: Caminho para o arquivo docx de entrada.
    :param texto: Texto a ser substituído (exemplo: "QR").
    :param imagem_path: Caminho da imagem a ser inserida no lugar do texto.
    :param output_path: Caminho para o arquivo de saída (docx).
    """
    # Abre o documento docx
    doc = Document(docx_path)

    # Itera pelas tabelas do documento
    for tabela in doc.tables:
        # Itera pelas linhas de cada tabela
        for linha in tabela.rows:
            # Itera pelas células de cada linha
            for celula in linha.cells:
                # Verifica se o texto da célula contém o texto a ser substituído
                if texto in celula.text:
                    # Substitui o texto na célula
                    for paragrafo in celula.paragraphs:
                        if texto in paragrafo.text:
                            # Remove o texto existente (se houver)
                            paragrafo.clear()

                            # Adiciona a imagem no parágrafo com tamanho 2x2 polegadas
                            run = paragrafo.add_run()
                            run.add_picture(imagem_path, width=Inches(2), height=Inches(2))
                            paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Salva o documento alterado
    doc.save(output_path)
    print(f"Documento salvo como {output_path}")

def obter_numero_do_arquivo(filename):
    """
    Obtém o número do arquivo a partir do nome do arquivo (assumindo que o número está no final do nome).
    :param filename: Nome do arquivo.
    :return: Número extraído do nome do arquivo.
    """
    base_nome = os.path.splitext(filename)[0]
    numero = ''.join(filter(str.isdigit, base_nome))  # Extrai apenas os números
    return int(numero)

if __name__ == "__main__":
    # Obtém o diretório onde o arquivo original está localizado
    script_dir = os.path.dirname(os.path.abspath(__file__))

    # Nome do arquivo de entrada (exemplo: "cartaodeembarque_semqrcode17.docx")
    base_filename = "cartaodeembarque_semqrcode14.docx"
    numero = obter_numero_do_arquivo(base_filename)

    # Caminhos completos para os arquivos no diretório do script
    docx_path = os.path.join(script_dir, base_filename)  # Arquivo de entrada
    imagem_path = os.path.join(script_dir, f"qrauto{numero:02d}.jpg")  # Imagem correspondente
    output_path = os.path.join(script_dir, f"cartao{numero:02d}.docx")  # Arquivo de saída

    # Chama a função para substituir o texto por imagem
    substituir_texto_por_imagem(docx_path, "QR", imagem_path, output_path)
