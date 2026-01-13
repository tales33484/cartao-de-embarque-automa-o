import docx
import re
import os
from datetime import datetime

# Lista de palavras que devem ser filtradas
palavras_proibidas = ["MR", "MRS", "MIS", "MST", "MISS", "MSTR"]

# Função para remover palavras proibidas de uma string
def filtrar_palavras(nome):
    palavras = nome.split()  # Divide o nome em palavras
    palavras_filtradas = [palavra for palavra in palavras if palavra.upper() not in palavras_proibidas]
    return " ".join(palavras_filtradas)

# Função para extrair a data do voo
def extrair_data_voo(linhas):
    meses = {
        "JAN": "01", "FEV": "02", "MAR": "03", "ABR": "04", "MAI": "05", "JUN": "06",
        "JUL": "07", "AGO": "08", "SET": "09", "OUT": "10", "NOV": "11", "DEZ": "12"
    }

    for linha in linhas:
        if linha.startswith("AP01"):
            data_match = re.search(r"\d{2}/[A-Za-z]{3}/\d{4}", linha)
            if data_match:
                data_str = data_match.group(0)
                dia, mes_abrev, ano = data_str.split("/")
                mes = meses.get(mes_abrev.upper(), "01")
                return f"{dia}/{mes}/{ano}"

    return None

# Função para extrair informações do passageiro
def extrair_informacoes_passageiro(arquivo, numero_passageiro):
    with open(arquivo, 'r', encoding='utf-8') as f:
        conteudo = f.readlines()

    # Extrair a data do voo
    data_voo = extrair_data_voo(conteudo)

    # Expressão regular para capturar as informações do passageiro
    pattern = r"(\d+)\s+([A-Z]{6})\s+([A-Za-z, ]+)"  # Captura o número, localizador e nome

    # Encontrar o passageiro baseado no número
    contador = 0
    for linha in conteudo:
        match = re.search(pattern, linha)
        if match:
            contador += 1
            if contador == numero_passageiro:  # Encontrou o passageiro com o número desejado
                numero, localizador, nome = match.groups()
                nomes = nome.split(", ")
                sobrenome = nomes[0]
                if len(nomes) > 1:
                    nomes_restantes = nomes[1].split(" ", 1)
                    primeiro_nome = nomes_restantes[0]
                    segundo_nome = nomes_restantes[1] if len(nomes_restantes) > 1 else ""
                else:
                    primeiro_nome, segundo_nome = "", ""
                sobrenome_filtrado = filtrar_palavras(sobrenome)
                segundo_nome_filtrado = filtrar_palavras(segundo_nome)
                return numero, localizador, primeiro_nome, segundo_nome_filtrado, sobrenome_filtrado, data_voo

    return None, None, None, None, None, None

# Função para substituir as informações no arquivo modelo.docx
def substituir_no_modelo(nome_arquivo, numero, localizador, primeiro_nome, segundo_nome, sobrenome, numero_passageiro, data_voo):
    # Carregar o arquivo DOCX
    doc = docx.Document(nome_arquivo)

    # Substituir "PAX" nos parágrafos
    for parágrafo in doc.paragraphs:
        if "PAX" in parágrafo.text:
            parágrafo.text = parágrafo.text.replace("PAX", f"{localizador} {primeiro_nome} {segundo_nome} {sobrenome}")

    # Substituir "PAX" nas tabelas
    for tabela in doc.tables:
        for linha in tabela.rows:
            for célula in linha.cells:
                if "PAX" in célula.text:
                    célula.text = célula.text.replace("PAX", f"{localizador} {primeiro_nome} {segundo_nome} {sobrenome}")

    # Substituir "DATAGERAL" pela data do voo
    for parágrafo in doc.paragraphs:
        if "DATAGERAL" in parágrafo.text:
            parágrafo.text = parágrafo.text.replace("DATAGERAL", data_voo)

    # Substituir "DATAGERAL" nas tabelas
    for tabela in doc.tables:
        for linha in tabela.rows:
            for célula in linha.cells:
                if "DATAGERAL" in célula.text:
                    célula.text = célula.text.replace("DATAGERAL", data_voo)

    # Salvar o documento modificado com o nome específico para o passageiro
    doc.save(os.path.join(current_folder, f'cartaodeembarque_semqrcode{str(numero_passageiro).zfill(2)}.docx'))
    print(f"Substituições realizadas e documento salvo como 'cartaodeembarque_semqrcode{str(numero_passageiro).zfill(2)}.docx'.")

# Obter o caminho do diretório onde o script está sendo executado
current_folder = os.path.dirname(os.path.abspath(__file__))

# Caminho absoluto do arquivo de entrada
arquivo_passageiros = os.path.join(current_folder, "conteudo_relatorio.txt")

# Caminho completo do arquivo modelo.docx
arquivo_modelo = os.path.join(current_folder, "modelo.docx")

# Definir o número do passageiro a ser processado
numero_passageiro = 13  # Modifique aqui o número do passageiro conforme necessário

# Extrair informações do passageiro
numero, localizador, primeiro_nome, segundo_nome, sobrenome, data_voo = extrair_informacoes_passageiro(arquivo_passageiros, numero_passageiro)

if numero and localizador and primeiro_nome and sobrenome and data_voo:
    # Substituir as informações no modelo
    substituir_no_modelo(arquivo_modelo, numero, localizador, primeiro_nome, segundo_nome, sobrenome, numero_passageiro, data_voo)
else:
    print(f"Erro: Passageiro {numero_passageiro} não encontrado ou data do voo não disponível no arquivo.")
